from langchain_community.document_loaders import PyPDFium2Loader, UnstructuredPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from docx import Document as DocxDocument
from logger import get_logger
import csv
import os
import subprocess
from pathlib import Path

log = get_logger(__name__)


def extract_pdf_images(
    file_path,
    user_id=None,
    upload_id=None,
    output_root="images",
    on_page_start=None,
    on_page_done=None,
):
    """
    Extracts PDF pages as PNG images using PAGE-SCOPED poppler subprocesses.

    Safety guarantees:
    - Each page is isolated in its own subprocess
    - One bad page does NOT fail the document
    - Native failures do NOT escape the subprocess
    - No exceptions are raised for per-page failures

    Observability guarantees:
    - Optional callbacks allow live progress reporting
    - Callbacks are best-effort and never affect extraction

    Primary storage:
    images/{user_id}/{upload_id}__{original_filename_stem}/page_{n}.png

    Returns a list of relative paths suitable for /images static serving.
    """

    filename_stem = Path(file_path).stem

    if upload_id:
        folder_name = f"{upload_id}__{filename_stem}"
    else:
        folder_name = filename_stem

    user_dir = os.path.join(output_root, str(user_id)) if user_id else output_root
    image_output_dir = os.path.join(user_dir, folder_name)
    os.makedirs(image_output_dir, exist_ok=True)

    log.info("pdf_image_conversion_started", file=file_path)

    saved_files = []

    try:
        loader = PyPDFium2Loader(file_path)
        docs = loader.load()
        total_pages = len(docs)
    except Exception as e:
        log.error("pdf_page_count_failed", file=file_path, error=str(e), exc_info=True)
        return []

    for page_num in range(1, total_pages + 1):
        if callable(on_page_start):
            try:
                on_page_start(page_num, total_pages)
            except Exception as e:
                log.warning("pdf_page_start_callback_failed", page=page_num, error=str(e))

        output_prefix = os.path.join(image_output_dir, f"page_{page_num}")

        cmd = [
            "pdftoppm",
            "-png",
            "-r",
            "150",
            "-f",
            str(page_num),
            "-l",
            str(page_num),
            file_path,
            output_prefix,
        ]

        try:
            subprocess.run(
                cmd,
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,  # per-page hard cap
            )

            num_digits = len(str(total_pages))
            expected_file = f"{output_prefix}-{page_num:0{num_digits}d}.png"

            if os.path.exists(expected_file):
                final_path = os.path.join(image_output_dir, f"page_{page_num}.png")
                os.rename(expected_file, final_path)
                saved_files.append(final_path.replace("\\", "/"))

                if callable(on_page_done):
                    try:
                        on_page_done(page_num, total_pages)
                    except Exception as e:
                        log.warning("pdf_page_done_callback_failed", page=page_num, error=str(e))
            else:
                log.warning("pdf_no_image_produced", page=page_num, file=file_path)

        except subprocess.TimeoutExpired:
            log.warning("pdf_image_page_timeout", page=page_num, file=file_path)
            continue
        except subprocess.CalledProcessError as e:
            log.warning("pdf_image_page_failed", page=page_num, file=file_path, error=e.stderr.decode(errors="ignore"))
            continue
        except Exception as e:
            log.warning("pdf_image_page_error", page=page_num, file=file_path, error=str(e))
            continue

    log.info("pdf_images_extracted", saved=len(saved_files), total_pages=total_pages, output_dir=image_output_dir)

    return saved_files


def load_pdf(file_path, user_id=None):
    try:
        loader = UnstructuredPDFLoader(file_path)
        docs = loader.load()
    except Exception as e:
        log.warning("pdf_unstructured_loader_failed", file=file_path, error=str(e))
        loader = PyPDFium2Loader(file_path)
        docs = loader.load()

    log.info("pdf_loaded", file=file_path, pages=len(docs))
    return docs


def load_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    metadata = {"source": file_path}
    return [Document(page_content=text, metadata=metadata)]


def load_csv(file_path):
    with open(file_path, newline='', encoding="utf-8") as csvfile:
        reader = csv.reader(csvfile)
        text = "\n".join([", ".join(row) for row in reader])
    metadata = {"source": file_path}
    return [Document(page_content=text, metadata=metadata)]


def load_docx(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    metadata = {"source": file_path}
    return [Document(page_content=text, metadata=metadata)]


def load_file(file_path, user_id=None):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".pdf":
        return load_pdf(file_path, user_id=user_id)
    elif ext == ".txt":
        return load_txt(file_path)
    elif ext == ".csv":
        return load_csv(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def chunk_documents(docs):
    """
    Split documents into chunks while preserving page_number metadata.
    Ensures every chunk has a page_number for correct image alignment.
    """
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    all_chunks = []

    for doc in docs:
        chunks = text_splitter.split_text(doc.page_content)

        page_number = doc.metadata.get("page_number") or doc.metadata.get("page") or None

        for i, chunk in enumerate(chunks):
            metadata = dict(doc.metadata)
            metadata["page_number"] = page_number if page_number is not None else i + 1
            all_chunks.append(Document(page_content=chunk, metadata=metadata))

    return all_chunks
